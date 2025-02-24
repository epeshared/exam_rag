from docx import Document
import os
import io
import re
from docx.shared import Inches
from log_config import get_logger
import json
from docx.image.exceptions import UnrecognizedImageError

with open("config.json", "r", encoding="utf-8") as f:
    config = json.load(f)

log_path = config.get("log_path")
logger = get_logger(__name__, log_file=log_path)


def process_chunk_content(doc, content):
    """
    处理文本内容，识别图片描述占位符并将其替换为实际图片。
    占位符格式假设为: [图片路径] 描述文本 [/]
    当成功插入图片后，跳过占位符内的描述文本（即删除这些文本）。
    """
    pattern = r'\[([^\]]+\.png)\](.*?)\[/\]'
    pos = 0
    for match in re.finditer(pattern, content, re.DOTALL):
        start, end = match.span()
        # 添加占位符之前的文本
        if start > pos:
            text_before = content[pos:start].strip()
            if text_before:
                doc.add_paragraph(text_before)
        image_path = match.group(1).strip()
        # 尝试插入图片
        try:
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(4))
            else:
                doc.add_paragraph(f"[图片 {image_path} 未找到]")
        except UnrecognizedImageError:
            doc.add_paragraph(f"[图片 {image_path} 格式无法识别]")
        # 无论占位符内描述内容是什么，都跳过，不添加到文档中
        pos = end
    # 添加剩余文本
    if pos < len(content):
        remaining = content[pos:].strip()
        if remaining:
            doc.add_paragraph(remaining)

def generate_search_docx(mistakes, config, course):
    """
    根据搜索结果生成 DOCX 文档。
    每个搜索结果的详细内容从对应的文本文件中读取，路径格式：
      os.path.join(config["base_chunks_output_dir"], course, doc_name, f"chunk_{chunk_id}.txt")
    返回 DOCX 文件的二进制内容（BytesIO 对象）。
    """
    doc = Document()
    doc.add_heading("相关错题汇总", level=1)
    
    base_chunks_dir = config["base_chunks_output_dir"]
    
    for item in mistakes:
        meta = item["metadata"]  
        doc_name = meta["doc"]
        chunk_id = meta["chunk_id"]
        distance = item.get("distance", 0)
        heading_text = f"文档: {doc_name} 题目编号: {chunk_id}, 距离: {distance:.4f}"
        doc.add_heading(heading_text, level=2)
        
        chunk_file_path = os.path.join(base_chunks_dir, course, doc_name, f"chunk_{chunk_id}.txt")
        if os.path.exists(chunk_file_path):
            with open(chunk_file_path, "r", encoding="utf-8") as f:
                chunk_content = f.read()
        else:
            chunk_content = "未找到题目内容文件。"
        process_chunk_content(doc, chunk_content)
        doc.add_page_break()
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
