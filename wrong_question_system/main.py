import os
import json
import numpy as np
import tempfile
import io
import streamlit as st
from PIL import Image
from docx import Document
from log_config import get_logger
from vdb_index import VectorDBIndex
from query_embedding import get_query_embedding_bge
from image_caption import load_chinese_image_captioning_model, chinese_image_caption_inference
from docx_generator import generate_search_docx

# 加载配置文件
with open("config.json", "r", encoding="utf-8") as f:
    config = json.load(f)

# 通过配置文件设置日志路径
log_path = config.get("log_path")
logger = get_logger(__name__, log_file=log_path)

st.title("错题整理系统 - 查询相关错题")
st.write("请输入查询文本或上传图片，系统将自动生成查询向量并检索相似错题，然后生成 DOCX 文档。")

# 文本输入框
text_input = st.text_area("查询文本（可选）", height=150)
# 图片上传组件
image_file = st.file_uploader("上传图片（可选）", type=["png", "jpg", "jpeg"])

# 初始化向量数据库索引（以“地理”课程为示例）
EMBEDDINGS_DIR = os.path.join(config["embedding_output_dir"], "地理")
vdb = VectorDBIndex(engine="faiss")
vdb.load_all_embeddings(EMBEDDINGS_DIR)
vdb.build_index()

# 加载图像描述模型（例如 Qwen2.5-VL-7B-Instruct）
VL_MODEL_PATH = config.get("vl_model_path", "/nvme0/models/Qwen/Qwen2.5-VL-7B-Instruct/")
processor, caption_model, caption_device = load_chinese_image_captioning_model(VL_MODEL_PATH)

# def generate_search_docx(results, config, course):
#     """
#     根据搜索结果生成 DOCX 文档。
#     每个搜索结果的详细内容从对应的文本文件中读取，路径格式：
#       os.path.join(config["base_chunks_output_dir"], course, doc_name, f"chunk_{chunk_id}.txt")
#     返回 DOCX 文件的二进制内容。
#     """
#     doc = Document()
#     doc.add_heading("相关错题汇总", level=1)
    
#     base_chunks_dir = config["base_chunks_output_dir"]
    
#     for res in results:
#         meta = res["metadata"]
#         distance = res["distance"]
#         doc_name = meta["doc"]
#         chunk_id = meta["chunk_id"]
#         # 添加标题信息
#         heading_text = f"文档: {doc_name} 题目编号: {chunk_id}, 距离: {distance:.4f}"
#         doc.add_heading(heading_text, level=2)
        
#         # 构造对应的文本文件路径
#         chunk_file_path = os.path.join(base_chunks_dir, course, doc_name, f"chunk_{chunk_id}.txt")
#         if os.path.exists(chunk_file_path):
#             with open(chunk_file_path, "r", encoding="utf-8") as f:
#                 chunk_content = f.read()
#         else:
#             chunk_content = "未找到题目内容文件。"
#         doc.add_paragraph(chunk_content)
#         doc.add_page_break()
    
#     # 将生成的 DOCX 保存到 BytesIO 对象中
#     file_stream = io.BytesIO()
#     doc.save(file_stream)
#     file_stream.seek(0)
#     return file_stream

def process_query(text_input, image_file):
    """
    根据用户输入的文本和上传的图片生成查询向量，
    利用向量索引搜索相似错题，并生成包含详细信息的 DOCX 文档供下载。
    """
    query_text = ""
    if image_file is not None:
        try:
            image = Image.open(image_file)
        except Exception as e:
            st.error("上传的图片无法打开，请检查格式。")
            logger.error("图片打开失败: %s", e, exc_info=True)
            return "图片无法打开。"
        # 保存图片到临时文件，以便调用图像描述模型
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            temp_path = tmp.name
            image.save(temp_path)
        caption = chinese_image_caption_inference(temp_path, processor, caption_model, caption_device)
        os.remove(temp_path)
        logger.info("图片描述生成: %s", caption)
        query_text = caption.strip()
    
    if text_input and text_input.strip():
        if query_text:
            query_text = text_input.strip() + " " + query_text
        else:
            query_text = text_input.strip()
    
    if not query_text:
        return "请输入查询文本或上传图片！"
    
    logger.info("最终查询文本：%s", query_text)
    # 使用 BGE 模型生成查询向量
    query_embedding = get_query_embedding_bge(query_text)
    if query_embedding is None:
        return "生成查询向量失败，请检查日志。"
    
    query_vector = np.array(query_embedding, dtype="float32").reshape(1, -1)
    results = vdb.search(query_vector, k=5)
    if not results:
        return "未找到相似题目。"
    
    # 生成 DOCX 文档
    docx_stream = generate_search_docx(results, config, "地理")
    
    # 使用 Streamlit 下载按钮提供下载
    st.download_button(
        label="下载相关错题文档",
        data=docx_stream,
        file_name="related_mistakes.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    return "相关错题文档已生成，请点击上方按钮下载。"

if st.button("查询相关错题"):
    result_message = process_query(text_input, image_file)
    st.text_area("提示信息", result_message, height=100)
